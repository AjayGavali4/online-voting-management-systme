"""
Screenshot Capture Script for Voting System
Captures all pages in both light and dark mode
"""

import os
import time
import subprocess
import sys
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
BASE_URL = "http://127.0.0.1:5000"
SCREENSHOT_DIR = "d:/ajay/screenshots"
OUTPUT_DOC = "d:/ajay/output.doc"

# Pages to capture (URL path, page name)
PAGES = [
    ("/", "Landing Page"),
    ("/auth/voter/login", "Voter Login"),
    ("/auth/voter/register", "Voter Registration"),
    ("/auth/admin/login", "Admin Login"),
    ("/auth/candidate/login", "Candidate Login"),
]

def setup_driver():
    """Setup Chrome driver with options"""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=chrome_options)
    return driver

def toggle_theme(driver, to_dark=True):
    """Toggle theme using the theme toggle button"""
    try:
        # Look for theme toggle button
        toggle_btn = driver.find_element(By.CSS_SELECTOR, ".theme-toggle, #theme-toggle, [data-theme-toggle]")
        if toggle_btn:
            driver.execute_script("arguments[0].click();", toggle_btn)
            time.sleep(0.5)
            return True
    except:
        pass
    
    # Fallback: directly toggle body class
    if to_dark:
        driver.execute_script("document.body.classList.add('dark-theme', 'dark-mode'); document.body.classList.remove('light-theme', 'light-mode');")
    else:
        driver.execute_script("document.body.classList.remove('dark-theme', 'dark-mode'); document.body.classList.add('light-theme', 'light-mode');")
    time.sleep(0.3)
    return True

def set_light_mode(driver):
    """Force light mode"""
    driver.execute_script("""
        document.body.classList.remove('dark-theme', 'dark-mode', 'dark-neon');
        document.body.classList.add('light-theme', 'light-mode');
        document.documentElement.setAttribute('data-theme', 'light');
        localStorage.setItem('theme', 'light');
    """)
    time.sleep(0.3)

def set_dark_mode(driver):
    """Force dark mode"""
    driver.execute_script("""
        document.body.classList.remove('light-theme', 'light-mode');
        document.body.classList.add('dark-theme', 'dark-mode', 'dark-neon');
        document.documentElement.setAttribute('data-theme', 'dark');
        localStorage.setItem('theme', 'dark');
    """)
    time.sleep(0.3)

def capture_page(driver, url, name, screenshot_dir):
    """Capture a page in both light and dark mode"""
    screenshots = {}
    
    print(f"Capturing: {name}")
    driver.get(url)
    time.sleep(1)  # Wait for page to load
    
    # Capture light mode
    set_light_mode(driver)
    time.sleep(0.5)
    light_path = os.path.join(screenshot_dir, f"{name.replace(' ', '_')}_light.png")
    driver.save_screenshot(light_path)
    screenshots['light'] = light_path
    print(f"  - Light mode saved: {light_path}")
    
    # Capture dark mode
    set_dark_mode(driver)
    time.sleep(0.5)
    dark_path = os.path.join(screenshot_dir, f"{name.replace(' ', '_')}_dark.png")
    driver.save_screenshot(dark_path)
    screenshots['dark'] = dark_path
    print(f"  - Dark mode saved: {dark_path}")
    
    return screenshots

def create_word_document(all_screenshots, output_path):
    """Create Word document with all screenshots"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Voting System - Output Pages', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This document contains screenshots of all system pages in both Light Mode and Dark Mode.")
    doc.add_paragraph()
    
    for page_name, screenshots in all_screenshots.items():
        # Page heading
        doc.add_heading(page_name, level=1)
        
        # Light Mode section
        doc.add_heading('Light Mode', level=2)
        if os.path.exists(screenshots['light']):
            doc.add_picture(screenshots['light'], width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Dark Mode section
        doc.add_heading('Dark Mode', level=2)
        if os.path.exists(screenshots['dark']):
            doc.add_picture(screenshots['dark'], width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break after each page
        doc.add_page_break()
    
    # Save
    # Save as .docx first, then we can copy to .doc
    docx_path = output_path.replace('.doc', '.docx') if output_path.endswith('.doc') else output_path
    doc.save(docx_path)
    print(f"\nDocument saved: {docx_path}")
    return docx_path

def main():
    # Create screenshot directory
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    
    print("=" * 50)
    print("Screenshot Capture Script")
    print("=" * 50)
    
    # Setup driver
    print("\nSetting up Chrome driver...")
    driver = setup_driver()
    
    all_screenshots = {}
    
    try:
        print(f"\nCapturing {len(PAGES)} pages...")
        print("-" * 50)
        
        for path, name in PAGES:
            url = BASE_URL + path
            screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
            all_screenshots[name] = screenshots
        
        print("-" * 50)
        print(f"\nTotal pages captured: {len(all_screenshots)}")
        
        # Create Word document
        print("\nCreating Word document...")
        output_file = create_word_document(all_screenshots, OUTPUT_DOC)
        
        print("\n" + "=" * 50)
        print("DONE! Screenshots captured and document created.")
        print(f"Document: {output_file}")
        print("=" * 50)
        
    finally:
        driver.quit()

if __name__ == "__main__":
    main()
