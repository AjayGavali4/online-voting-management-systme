"""
Complete Screenshot Capture - With Full Login Flow
Handles OTP verification for voter and candidate
"""

import os
import time
import re
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
BASE_URL = "http://127.0.0.1:5001"
SCREENSHOT_DIR = "d:/ajay/output"
OUTPUT_DOC = "d:/ajay/output.doc"

# ============ ALL PAGES ============

AUTH_PAGES = [
    ("/", "01_Landing_Page"),
    ("/auth/voter/login", "02_Voter_Login_Page"),
    ("/auth/voter/register", "03_Voter_Registration_Page"),
    ("/auth/admin/login", "04_Admin_Login_Page"),
    ("/auth/candidate/login", "05_Candidate_Login_Page"),
    ("/auth/candidate/register", "06_Candidate_Registration_Page"),
]

ADMIN_PAGES = [
    ("/admin/dashboard", "07_Admin_Dashboard"),
    ("/admin/voters", "08_Admin_Voters_List"),
    ("/admin/voter/create", "09_Admin_Create_Voter"),
    ("/admin/candidates", "10_Admin_Candidates_List"),
    ("/admin/candidate/create", "11_Admin_Create_Candidate"),
    ("/admin/documents/pending", "12_Admin_Pending_Docs"),
    ("/admin/documents", "13_Admin_All_Documents"),
    ("/admin/elections", "14_Admin_Elections_List"),
    ("/admin/election/create", "15_Admin_Create_Election"),
    ("/admin/audit-logs", "16_Admin_Audit_Logs"),
    ("/admin/fraud-detection", "17_Admin_Fraud_Detection"),
    ("/admin/settings/2fa", "18_Admin_2FA_Settings"),
]

VOTER_PAGES = [
    ("/voter/dashboard", "19_Voter_Dashboard"),
    ("/voter/elections", "20_Voter_Elections_List"),
    ("/voter/documents", "21_Voter_My_Documents"),
    ("/voter/upload-document", "22_Voter_Upload_Document"),
    ("/voter/my-votes", "23_Voter_My_Votes_History"),
    ("/voter/profile", "24_Voter_Edit_Profile"),
    ("/voter/notifications", "25_Voter_Notifications"),
    ("/voter/help-support", "26_Voter_Help_Support"),
    ("/voter/verification-status", "27_Voter_Verification_Status"),
]

CANDIDATE_PAGES = [
    ("/candidate/dashboard", "28_Candidate_Dashboard"),
    ("/candidate/elections", "29_Candidate_My_Elections"),
    ("/candidate/available-elections", "30_Candidate_Available_Elections"),
    ("/candidate/profile", "31_Candidate_Edit_Profile"),
    ("/candidate/statistics", "32_Candidate_Statistics"),
]

def setup_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1200")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=chrome_options)

def set_light_mode(driver):
    driver.execute_script("""
        document.body.classList.remove('dark-theme', 'dark-mode', 'dark-neon');
        document.body.classList.add('light-theme', 'light-mode');
        document.documentElement.setAttribute('data-theme', 'light');
        localStorage.setItem('theme', 'light');
    """)
    time.sleep(0.3)

def set_dark_mode(driver):
    driver.execute_script("""
        document.body.classList.remove('light-theme', 'light-mode');
        document.body.classList.add('dark-theme', 'dark-mode', 'dark-neon');
        document.documentElement.setAttribute('data-theme', 'dark');
        localStorage.setItem('theme', 'dark');
    """)
    time.sleep(0.3)

def capture_page(driver, url, name, screenshot_dir):
    screenshots = {}
    print(f"  {name}")
    try:
        driver.get(url)
        time.sleep(1.5)
        
        if "/login" in driver.current_url and "/login" not in url:
            print(f"    [Skipped - redirected to login]")
            return None
        
        set_light_mode(driver)
        time.sleep(0.5)
        light_path = os.path.join(screenshot_dir, f"{name}_light.png")
        driver.save_screenshot(light_path)
        screenshots['light'] = light_path
        
        set_dark_mode(driver)
        time.sleep(0.5)
        dark_path = os.path.join(screenshot_dir, f"{name}_dark.png")
        driver.save_screenshot(dark_path)
        screenshots['dark'] = dark_path
        
        print(f"    [OK]")
        return screenshots
    except Exception as e:
        print(f"    [Error: {str(e)[:40]}]")
        return None

def login_admin(driver):
    """Login as admin (no OTP required)"""
    print("\n[Admin Login]")
    driver.get(BASE_URL + "/auth/admin/login")
    time.sleep(1)
    
    try:
        driver.find_element(By.NAME, "username").send_keys("admin")
        driver.find_element(By.NAME, "password").send_keys("admin123")
        submit = driver.find_element(By.CSS_SELECTOR, "button[type='submit']")
        driver.execute_script("arguments[0].click();", submit)
        time.sleep(2)
        
        if "/admin/" in driver.current_url:
            print("  Success!")
            return True
        print(f"  Current URL: {driver.current_url}")
        return True
    except Exception as e:
        print(f"  Error: {e}")
        return False

def login_voter_direct(driver):
    """Login voter using test route (bypasses OTP)"""
    print("\n[Voter Login - Direct Test Route]")
    
    driver.get(BASE_URL + "/test-login/voter")
    time.sleep(2)
    
    if "/voter/" in driver.current_url:
        print("  Success! Logged in as voter")
        return True
    else:
        print(f"  Result URL: {driver.current_url}")
        return False

def login_candidate_direct(driver):
    """Login candidate using test route (bypasses OTP)"""
    print("\n[Candidate Login - Direct Test Route]")
    
    driver.get(BASE_URL + "/test-login/candidate")
    time.sleep(2)
    
    if "/candidate/" in driver.current_url:
        print("  Success! Logged in as candidate")
        return True
    else:
        print(f"  Result URL: {driver.current_url}")
        return False

def logout(driver):
    try:
        driver.get(BASE_URL + "/auth/logout")
        time.sleep(1)
    except:
        pass

def create_word_document(all_screenshots, output_path):
    doc = Document()
    
    title = doc.add_heading('Voting System - Complete Output Pages', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total Pages: {len(all_screenshots)}")
    doc.add_paragraph("All screenshots in Light Mode and Dark Mode")
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for i, page_name in enumerate(all_screenshots.keys(), 1):
        clean_name = page_name.replace('_', ' ')
        doc.add_paragraph(f"{i}. {clean_name}")
    doc.add_page_break()
    
    # Screenshots
    for page_name, screenshots in all_screenshots.items():
        clean_name = page_name.replace('_', ' ')
        doc.add_heading(clean_name, level=1)
        
        doc.add_heading('Light Mode', level=2)
        if screenshots and 'light' in screenshots and os.path.exists(screenshots['light']):
            doc.add_picture(screenshots['light'], width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        doc.add_heading('Dark Mode', level=2)
        if screenshots and 'dark' in screenshots and os.path.exists(screenshots['dark']):
            doc.add_picture(screenshots['dark'], width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    doc.save(output_path)
    print(f"\nDocument saved: {output_path}")
    return output_path

def main():
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    
    print("=" * 60)
    print("  COMPLETE SCREENSHOT CAPTURE")
    print("=" * 60)
    
    driver = setup_driver()
    all_screenshots = {}
    
    try:
        # ===== SECTION 1: AUTH PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 1: AUTH & PUBLIC PAGES")
        print("=" * 60)
        
        for path, name in AUTH_PAGES:
            url = BASE_URL + path
            screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
            if screenshots:
                all_screenshots[name] = screenshots
        
        # ===== SECTION 2: ADMIN PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 2: ADMIN PAGES (All Functions)")
        print("=" * 60)
        
        if login_admin(driver):
            for path, name in ADMIN_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        
        # ===== SECTION 3: VOTER PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 3: VOTER PAGES (All Functions)")
        print("=" * 60)
        
        if login_voter_direct(driver):
            for path, name in VOTER_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        else:
            print("  Voter login failed - skipping voter pages")
        
        # ===== SECTION 4: CANDIDATE PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 4: CANDIDATE PAGES (All Functions)")
        print("=" * 60)
        
        if login_candidate_direct(driver):
            for path, name in CANDIDATE_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        else:
            print("  Candidate login failed - skipping candidate pages")
        
        # ===== SECTION 5: OTP PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 5: OTP VERIFICATION PAGES")
        print("=" * 60)
        
        # Capture Voter Login with OTP section
        driver.get(BASE_URL + "/auth/voter/login")
        time.sleep(1)
        otp_screenshots = capture_page(driver, driver.current_url, "33_Voter_Login_OTP_Page", SCREENSHOT_DIR)
        if otp_screenshots:
            all_screenshots["33_Voter_Login_OTP_Page"] = otp_screenshots
        
        # Capture Candidate Login with OTP section
        driver.get(BASE_URL + "/auth/candidate/login")
        time.sleep(1)
        otp_screenshots = capture_page(driver, driver.current_url, "34_Candidate_Login_OTP_Page", SCREENSHOT_DIR)
        if otp_screenshots:
            all_screenshots["34_Candidate_Login_OTP_Page"] = otp_screenshots
        
        # ===== CREATE DOCUMENT =====
        print("\n" + "=" * 60)
        print(f"TOTAL PAGES: {len(all_screenshots)}")
        print("Creating Word document...")
        
        create_word_document(all_screenshots, OUTPUT_DOC)
        
        print("\n" + "=" * 60)
        print("  COMPLETE!")
        print(f"  Screenshots: {SCREENSHOT_DIR}")
        print(f"  Document: {OUTPUT_DOC}")
        print("=" * 60)
        
    finally:
        driver.quit()

if __name__ == "__main__":
    main()
