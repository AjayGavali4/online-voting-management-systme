"""
Complete Screenshot Capture Script for Voting System
Captures ALL pages including OTP, Audit Logs, 2FA for Admin, Voter, Candidate
"""

import os
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
BASE_URL = "http://127.0.0.1:5001"
SCREENSHOT_DIR = "d:/ajay/output"
OUTPUT_DOCX = "d:/ajay/output.doc"

# ============ ALL PAGES TO CAPTURE ============

# 1. PUBLIC/AUTH PAGES
AUTH_PAGES = [
    ("/", "01_Landing_Page"),
    ("/auth/voter/login", "02_Voter_Login"),
    ("/auth/voter/register", "03_Voter_Register"),
    ("/auth/admin/login", "04_Admin_Login"),
    ("/auth/candidate/login", "05_Candidate_Login"),
    ("/auth/candidate/register", "06_Candidate_Register"),
]

# 2. ADMIN PAGES (after login) - ALL FUNCTIONS
ADMIN_PAGES = [
    ("/admin/dashboard", "07_Admin_Dashboard"),
    ("/admin/voters", "08_Admin_Voters_List"),
    ("/admin/voter/create", "09_Admin_Create_Voter"),
    ("/admin/candidates", "10_Admin_Candidates_List"),
    ("/admin/candidate/create", "11_Admin_Create_Candidate"),
    ("/admin/documents/pending", "12_Admin_Pending_Documents"),
    ("/admin/documents", "13_Admin_All_Documents"),
    ("/admin/elections", "14_Admin_Elections_List"),
    ("/admin/election/create", "15_Admin_Create_Election"),
    ("/admin/audit-logs", "16_Admin_Audit_Logs"),
    ("/admin/fraud-detection", "17_Admin_Fraud_Detection"),
    ("/admin/settings/2fa", "18_Admin_2FA_Settings"),
]

# 3. VOTER PAGES (after login) - ALL FUNCTIONS
VOTER_PAGES = [
    ("/voter/dashboard", "19_Voter_Dashboard"),
    ("/voter/elections", "20_Voter_Elections_List"),
    ("/voter/documents", "21_Voter_My_Documents"),
    ("/voter/upload-document", "22_Voter_Upload_Document"),
    ("/voter/my-votes", "23_Voter_My_Votes"),
    ("/voter/profile", "24_Voter_Edit_Profile"),
    ("/voter/notifications", "25_Voter_Notifications"),
    ("/voter/help-support", "26_Voter_Help_Support"),
    ("/voter/verification-status", "27_Voter_Verification_Status"),
]

# 4. CANDIDATE PAGES (after login) - ALL FUNCTIONS
CANDIDATE_PAGES = [
    ("/candidate/dashboard", "28_Candidate_Dashboard"),
    ("/candidate/elections", "29_Candidate_My_Elections"),
    ("/candidate/available-elections", "30_Candidate_Available_Elections"),
    ("/candidate/profile", "31_Candidate_Edit_Profile"),
    ("/candidate/statistics", "32_Candidate_Statistics"),
]

def setup_driver():
    """Setup Chrome driver"""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1200")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=chrome_options)
    return driver

def set_light_mode(driver):
    """Force light mode"""
    driver.execute_script("""
        document.body.classList.remove('dark-theme', 'dark-mode', 'dark-neon');
        document.body.classList.add('light-theme', 'light-mode');
        document.documentElement.setAttribute('data-theme', 'light');
        localStorage.setItem('theme', 'light');
    """)
    time.sleep(0.3)

def set_dark_mode(driver):
    """Force dark mode"""
    driver.execute_script("""
        document.body.classList.remove('light-theme', 'light-mode');
        document.body.classList.add('dark-theme', 'dark-mode', 'dark-neon');
        document.documentElement.setAttribute('data-theme', 'dark');
        localStorage.setItem('theme', 'dark');
    """)
    time.sleep(0.3)

def capture_page(driver, url, name, screenshot_dir):
    """Capture a page in both light and dark mode"""
    screenshots = {}
    
    print(f"  {name}")
    try:
        driver.get(url)
        time.sleep(1.5)
        
        # Check if redirected to login
        if "/login" in driver.current_url and "/login" not in url:
            print(f"    [Skipped - redirected]")
            return None
        
        # Light mode
        set_light_mode(driver)
        time.sleep(0.5)
        light_path = os.path.join(screenshot_dir, f"{name}_light.png")
        driver.save_screenshot(light_path)
        screenshots['light'] = light_path
        
        # Dark mode
        set_dark_mode(driver)
        time.sleep(0.5)
        dark_path = os.path.join(screenshot_dir, f"{name}_dark.png")
        driver.save_screenshot(dark_path)
        screenshots['dark'] = dark_path
        
        print(f"    [OK]")
        return screenshots
    except Exception as e:
        print(f"    [Error: {str(e)[:40]}]")
        return None

def capture_otp_page(driver, name, screenshot_dir):
    """Capture OTP verification page with multiple OTP inputs"""
    screenshots = {}
    print(f"  {name}")
    
    try:
        # Light mode
        set_light_mode(driver)
        time.sleep(0.5)
        light_path = os.path.join(screenshot_dir, f"{name}_light.png")
        driver.save_screenshot(light_path)
        screenshots['light'] = light_path
        
        # Dark mode
        set_dark_mode(driver)
        time.sleep(0.5)
        dark_path = os.path.join(screenshot_dir, f"{name}_dark.png")
        driver.save_screenshot(dark_path)
        screenshots['dark'] = dark_path
        
        print(f"    [OK]")
        return screenshots
    except Exception as e:
        print(f"    [Error]")
        return None

def login_admin(driver):
    """Login as admin"""
    print("\n[Logging in as Admin]")
    driver.get(BASE_URL + "/auth/admin/login")
    time.sleep(1)
    
    try:
        username = driver.find_element(By.NAME, "username")
        password = driver.find_element(By.NAME, "password")
        username.clear()
        username.send_keys("admin")
        password.clear()
        password.send_keys("admin123")
        
        submit = driver.find_element(By.CSS_SELECTOR, "button[type='submit'], input[type='submit']")
        driver.execute_script("arguments[0].click();", submit)
        time.sleep(2)
        
        print("  Admin login successful")
        return True
    except Exception as e:
        print(f"  Admin login failed: {e}")
        return False

def login_voter(driver):
    """Login as voter"""
    print("\n[Logging in as Voter]")
    driver.get(BASE_URL + "/auth/voter/login")
    time.sleep(1)
    
    try:
        username = driver.find_element(By.NAME, "username")
        password = driver.find_element(By.NAME, "password")
        username.clear()
        username.send_keys("testvoter")
        password.clear()
        password.send_keys("voter123")
        
        submit = driver.find_element(By.CSS_SELECTOR, "button[type='submit'], input[type='submit']")
        driver.execute_script("arguments[0].click();", submit)
        time.sleep(2)
        
        print("  Voter login successful")
        return True
    except Exception as e:
        print(f"  Voter login failed: {e}")
        return False

def login_candidate(driver):
    """Login as candidate"""
    print("\n[Logging in as Candidate]")
    driver.get(BASE_URL + "/auth/candidate/login")
    time.sleep(1)
    
    try:
        username = driver.find_element(By.NAME, "username")
        password = driver.find_element(By.NAME, "password")
        username.clear()
        username.send_keys("testcandidate")
        password.clear()
        password.send_keys("candidate123")
        
        submit = driver.find_element(By.CSS_SELECTOR, "button[type='submit'], input[type='submit']")
        driver.execute_script("arguments[0].scrollIntoView(true);", submit)
        time.sleep(0.3)
        driver.execute_script("arguments[0].click();", submit)
        time.sleep(2)
        
        print("  Candidate login successful")
        return True
    except Exception as e:
        print(f"  Candidate login failed: {e}")
        return False

def logout(driver):
    """Logout"""
    try:
        driver.get(BASE_URL + "/auth/logout")
        time.sleep(1)
    except:
        pass

def trigger_otp_page(driver, user_type, screenshot_dir):
    """Trigger OTP verification and capture the page"""
    print(f"\n[Capturing OTP Page for {user_type}]")
    
    # Navigate to login page with OTP option
    if user_type == "voter":
        driver.get(BASE_URL + "/auth/voter/login")
    elif user_type == "candidate":
        driver.get(BASE_URL + "/auth/candidate/login")
    else:
        return None
    
    time.sleep(1)
    
    # Try to find and click OTP login tab/button
    try:
        otp_tab = driver.find_element(By.CSS_SELECTOR, "[data-tab='otp'], .otp-tab, #otp-login-tab, a[href*='otp']")
        driver.execute_script("arguments[0].click();", otp_tab)
        time.sleep(1)
    except:
        pass
    
    # If OTP modal/section is visible, capture it
    name = f"32_OTP_Verification_{user_type.capitalize()}"
    return capture_otp_page(driver, name, screenshot_dir)

def create_word_document(all_screenshots, output_path):
    """Create Word document with all screenshots"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Voting System - All Output Pages', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Complete screenshots of all system pages in Light Mode and Dark Mode.")
    doc.add_paragraph(f"Total Pages: {len(all_screenshots)}")
    doc.add_paragraph()
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    for i, page_name in enumerate(all_screenshots.keys(), 1):
        clean_name = page_name.replace('_', ' ')
        doc.add_paragraph(f"{i}. {clean_name}")
    doc.add_page_break()
    
    # Add screenshots
    for page_name, screenshots in all_screenshots.items():
        clean_name = page_name.replace('_', ' ')
        doc.add_heading(clean_name, level=1)
        
        # Light Mode
        doc.add_heading('Light Mode', level=2)
        if screenshots and 'light' in screenshots and os.path.exists(screenshots['light']):
            doc.add_picture(screenshots['light'], width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Dark Mode
        doc.add_heading('Dark Mode', level=2)
        if screenshots and 'dark' in screenshots and os.path.exists(screenshots['dark']):
            doc.add_picture(screenshots['dark'], width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    doc.save(output_path)
    print(f"\nDocument saved: {output_path}")
    return output_path

def main():
    # Create output directory
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    
    print("=" * 60)
    print("  VOTING SYSTEM - COMPLETE SCREENSHOT CAPTURE")
    print("=" * 60)
    
    driver = setup_driver()
    all_screenshots = {}
    
    try:
        # ===== SECTION 1: AUTH/PUBLIC PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 1: AUTH & PUBLIC PAGES")
        print("=" * 60)
        
        for path, name in AUTH_PAGES:
            url = BASE_URL + path
            screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
            if screenshots:
                all_screenshots[name] = screenshots
        
        # ===== SECTION 2: ADMIN PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 2: ADMIN PAGES")
        print("=" * 60)
        
        if login_admin(driver):
            for path, name in ADMIN_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        
        # ===== SECTION 3: VOTER PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 3: VOTER PAGES")
        print("=" * 60)
        
        if login_voter(driver):
            for path, name in VOTER_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        
        # ===== SECTION 4: CANDIDATE PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 4: CANDIDATE PAGES")
        print("=" * 60)
        
        if login_candidate(driver):
            for path, name in CANDIDATE_PAGES:
                url = BASE_URL + path
                screenshots = capture_page(driver, url, name, SCREENSHOT_DIR)
                if screenshots:
                    all_screenshots[name] = screenshots
            logout(driver)
        
        # ===== SECTION 5: OTP PAGES =====
        print("\n" + "=" * 60)
        print("SECTION 5: OTP VERIFICATION PAGES")
        print("=" * 60)
        
        # Capture Voter OTP Login Page
        driver.get(BASE_URL + "/auth/voter/login")
        time.sleep(1)
        try:
            otp_tabs = driver.find_elements(By.CSS_SELECTOR, "[data-tab='otp'], .otp-tab, #otp-tab, .tab-otp, [onclick*='otp'], button[class*='otp'], a[class*='otp']")
            for tab in otp_tabs:
                driver.execute_script("arguments[0].click();", tab)
                time.sleep(0.5)
                break
        except:
            pass
        otp_screenshots = capture_page(driver, driver.current_url, "33_Voter_OTP_Login", SCREENSHOT_DIR)
        if otp_screenshots:
            all_screenshots["33_Voter_OTP_Login"] = otp_screenshots
        
        # Capture Candidate OTP Login Page
        driver.get(BASE_URL + "/auth/candidate/login")
        time.sleep(1)
        try:
            otp_tabs = driver.find_elements(By.CSS_SELECTOR, "[data-tab='otp'], .otp-tab, #otp-tab, .tab-otp, [onclick*='otp'], button[class*='otp'], a[class*='otp']")
            for tab in otp_tabs:
                driver.execute_script("arguments[0].click();", tab)
                time.sleep(0.5)
                break
        except:
            pass
        otp_screenshots = capture_page(driver, driver.current_url, "34_Candidate_OTP_Login", SCREENSHOT_DIR)
        if otp_screenshots:
            all_screenshots["34_Candidate_OTP_Login"] = otp_screenshots
        
        # Capture Admin OTP/2FA Page
        driver.get(BASE_URL + "/auth/admin/login")
        time.sleep(1)
        try:
            otp_tabs = driver.find_elements(By.CSS_SELECTOR, "[data-tab='otp'], .otp-tab, #otp-tab, .tab-otp, [onclick*='otp'], button[class*='otp'], a[class*='otp']")
            for tab in otp_tabs:
                driver.execute_script("arguments[0].click();", tab)
                time.sleep(0.5)
                break
        except:
            pass
        otp_screenshots = capture_page(driver, driver.current_url, "35_Admin_OTP_Login", SCREENSHOT_DIR)
        if otp_screenshots:
            all_screenshots["35_Admin_OTP_Login"] = otp_screenshots
        
        # ===== CREATE DOCUMENT =====
        print("\n" + "=" * 60)
        print(f"TOTAL PAGES CAPTURED: {len(all_screenshots)}")
        print("Creating Word document...")
        
        output_file = create_word_document(all_screenshots, OUTPUT_DOCX)
        
        print("\n" + "=" * 60)
        print("  COMPLETE!")
        print(f"  Screenshots: {SCREENSHOT_DIR}")
        print(f"  Document: {output_file}")
        print("=" * 60)
        
    finally:
        driver.quit()

if __name__ == "__main__":
    main()
